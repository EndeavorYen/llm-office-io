#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Unit Tests for Word Editor New Features (v1.3.0)
Testing: add_image, insert_table, update_table_cell, set_paragraph_format, add_page_break
"""

import unittest
import os
import tempfile
import shutil
from pathlib import Path

# Add parent directory to path
import sys
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from src.word_editor import WordEditor
from docx import Document


class TestWordEditorNewFeatures(unittest.TestCase):
    """測試 Word Editor v1.3.0 新功能"""
    
    @classmethod
    def setUpClass(cls):
        """設置測試環境"""
        cls.test_dir = tempfile.mkdtemp()
        cls.test_docx = os.path.join(cls.test_dir, "test.docx")
        
        # 創建測試文檔
        doc = Document()
        doc.add_heading("測試文檔", 0)
        doc.add_paragraph("這是第一段內容。")
        doc.add_paragraph("第一章")
        doc.add_paragraph("這是第一章的內容。")
        doc.save(cls.test_docx)
    
    @classmethod
    def tearDownClass(cls):
        """清理測試環境"""
        shutil.rmtree(cls.test_dir)
    
    def setUp(self):
        """每個測試前的準備"""
        self.output_file = os.path.join(self.test_dir, f"output_{self._testMethodName}.docx")
        shutil.copy(self.test_docx, self.output_file)
        self.editor = WordEditor(self.output_file)
    
    def test_add_image_to_end(self):
        """測試在文檔末尾插入圖片"""
        # 創建測試圖片
        test_image = os.path.join(self.test_dir, "test.png")
        from PIL import Image
        img = Image.new('RGB', (100, 100), color='red')
        img.save(test_image)
        
        # 插入圖片
        result = self.editor.add_image(test_image, width_cm=5.0)
        self.assertTrue(result)
        
        # 儲存並驗證
        self.editor.save()
        doc = Document(self.output_file)
        
        # 檢查是否有內嵌圖片
        has_image = False
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if hasattr(run, '_element') and run._element.findall('.//{*}drawing'):
                    has_image = True
                    break
        
        self.assertTrue(has_image, "文檔應包含插入的圖片")
    
    def test_add_image_with_position(self):
        """測試在特定位置後插入圖片"""
        test_image = os.path.join(self.test_dir, "test.png")
        from PIL import Image
        img = Image.new('RGB', (100, 100), color='blue')
        img.save(test_image)
        
        result = self.editor.add_image(test_image, width_cm=3.0, position="第一章")
        self.assertTrue(result)
    
    def test_add_image_file_not_found(self):
        """測試圖片檔案不存在的情況"""
        result = self.editor.add_image("nonexistent.png")
        self.assertFalse(result)
    
    def test_insert_table_basic(self):
        """測試插入基本表格"""
        result = self.editor.insert_table(rows=3, cols=4)
        self.assertTrue(result)
        
        self.editor.save()
        doc = Document(self.output_file)
        
        # 驗證表格被插入
        self.assertGreater(len(doc.tables), 0, "應該至少有一個表格")
        table = doc.tables[-1]  # 最後一個表格
        self.assertEqual(len(table.rows), 3)
        self.assertEqual(len(table.columns), 4)
    
    def test_insert_table_with_data(self):
        """測試插入帶資料的表格"""
        data = [
            ["姓名", "年齡", "城市"],
            ["張三", "25", "台北"],
            ["李四", "30", "高雄"]
        ]
        
        result = self.editor.insert_table(rows=3, cols=3, data=data)
        self.assertTrue(result)
        
        self.editor.save()
        doc = Document(self.output_file)
        table = doc.tables[-1]
        
        # 驗證資料
        self.assertEqual(table.rows[0].cells[0].text, "姓名")
        self.assertEqual(table.rows[1].cells[1].text, "25")
    
    def test_insert_table_invalid_size(self):
        """測試無效的表格大小"""
        result = self.editor.insert_table(rows=0, cols=3)
        self.assertFalse(result)
    
    def test_update_table_cell(self):
        """測試更新表格儲存格"""
        # 先插入表格
        data = [["A", "B"], ["C", "D"]]
        self.editor.insert_table(rows=2, cols=2, data=data)
        self.editor.save()
        
        # 重新載入並更新
        editor2 = WordEditor(self.output_file)
        result = editor2.update_table_cell(0, 0, 1, "Updated")
        self.assertTrue(result)
        
        editor2.save()
        doc = Document(self.output_file)
        self.assertEqual(doc.tables[0].rows[0].cells[1].text, "Updated")
    
    def test_update_table_cell_invalid_index(self):
        """測試無效的表格索引"""
        result = self.editor.update_table_cell(99, 0, 0, "text")
        self.assertFalse(result)
    
    def test_set_paragraph_format_basic(self):
        """測試設定段落格式"""
        result = self.editor.set_paragraph_format(
            search_text="第一段",
            font_size=14,
            bold=True
        )
        self.assertTrue(result)
        
        self.editor.save()
        doc = Document(self.output_file)
        
        # 找到包含「第一段」的段落
        for para in doc.paragraphs:
            if "第一段" in para.text:
                for run in para.runs:
                    self.assertTrue(run.font.bold)
                break
    
    def test_set_paragraph_format_with_alignment(self):
        """測試設定對齊方式"""
        result = self.editor.set_paragraph_format(
            search_text="測試文檔",
            alignment="center"
        )
        self.assertTrue(result)
    
    def test_set_paragraph_format_not_found(self):
        """測試搜尋文字不存在"""
        result = self.editor.set_paragraph_format(
            search_text="不存在的文字",
            bold=True
        )
        self.assertFalse(result)
    
    def test_add_page_break_at_end(self):
        """測試在末尾插入分頁符號"""
        result = self.editor.add_page_break()
        self.assertTrue(result)
        
        self.editor.save()
        # 分頁符號已插入（無法直接驗證，但確保無錯誤）
    
    def test_add_page_break_after_text(self):
        """測試在特定文字後插入分頁符號"""
        result = self.editor.add_page_break(after_text="第一章")
        self.assertTrue(result)
    
    def test_add_page_break_text_not_found(self):
        """測試文字不存在時插入分頁"""
        result = self.editor.add_page_break(after_text="不存在")
        self.assertFalse(result)


class TestWordEditorIntegration(unittest.TestCase):
    """整合測試：組合多個功能"""
    
    def setUp(self):
        """準備測試環境"""
        self.test_dir = tempfile.mkdtemp()
        self.test_file = os.path.join(self.test_dir, "integration_test.docx")
    
    def tearDown(self):
        """清理"""
        shutil.rmtree(self.test_dir)
    
    def test_complete_workflow(self):
        """測試完整工作流程"""
        # 創建新文檔
        doc = Document()
        doc.add_heading("年度報告", 0)
        doc.add_paragraph("執行摘要")
        doc.save(self.test_file)
        
        # 使用編輯器
        editor = WordEditor(self.test_file)
        
        # 1. 設定標題格式
        editor.set_paragraph_format("年度報告", font_size=24, bold=True, alignment="center")
        
        # 2. 插入表格
        data = [
            ["項目", "數值"],
            ["營收", "100M"],
            ["成本", "60M"]
        ]
        editor.insert_table(2, 2, data=data, position="執行摘要")
        
        # 3. 添加分頁
        editor.add_page_break(after_text="執行摘要")
        
        # 4. 儲存
        editor.save()
        
        # 驗證
        doc = Document(self.test_file)
        self.assertGreater(len(doc.tables), 0)


if __name__ == '__main__':
    # 檢查 Pillow 是否安裝
    try:
        from PIL import Image
    except ImportError:
        print("警告: 需要安裝 Pillow 來運行圖片測試")
        print("請執行: pip install Pillow")
    
    unittest.main(verbosity=2)
