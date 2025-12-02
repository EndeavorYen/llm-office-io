#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Interactive Word Document Editor
直接操作 Word 文檔的互動式編輯工具
"""

from typing import Optional, List
import os
import sys
import argparse

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .constants import (
    MAX_PREVIEW_LENGTH,
    MAX_TEXT_DISPLAY,
    SUCCESS_SYMBOL,
    ERROR_SYMBOL,
    WORD_EXTENSION,
    DEFAULT_HEADING_LEVEL
)


class WordEditor:
    """Word 文檔編輯器類"""
    
    def __init__(self, filepath: str) -> None:
        """初始化 Word 編輯器
        
        Args:
            filepath: Word 文檔路徑
            
        Raises:
            FileNotFoundError: 當檔案不存在時
            ValueError: 當檔案格式不支援時
            RuntimeError: 當無法開啟文檔時
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"檔案不存在: {filepath}")
        
        if not filepath.endswith(WORD_EXTENSION):
            raise ValueError(f"不支援的檔案格式，需要 {WORD_EXTENSION}: {filepath}")
        
        try:
            self.filepath = filepath
            self.doc = Document(filepath)
        except Exception as e:
            raise RuntimeError(f"無法開啟文檔: {e}") from e
    
    def save(self, output_path: Optional[str] = None) -> None:
        """儲存文檔
        
        Args:
            output_path: 輸出路徑，None 表示覆蓋原檔案
        """
        save_path = output_path or self.filepath
        try:
            self.doc.save(save_path)
            print(f"{SUCCESS_SYMBOL} 文檔已儲存: {save_path}")
        except Exception as e:
            print(f"{ERROR_SYMBOL} 儲存失敗: {e}")
            raise
    
    def list_structure(self) -> None:
        """列出文檔結構（標題和段落）"""
        print("\n=== 文檔結構 ===\n")
        for i, para in enumerate(self.doc.paragraphs):
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                indent = "  " * (int(level) - 1) if level.isdigit() else ""
                print(f"[{i}] {indent}📌 {para.text[:MAX_PREVIEW_LENGTH]}")
            elif para.text.strip():
                preview = para.text[:MAX_PREVIEW_LENGTH].replace('\n', ' ')
                print(f"[{i}]    {preview}")
        print()
    
    def add_paragraph_after(
        self, 
        search_text: str, 
        new_content: str, 
        heading_level: Optional[int] = None
    ) -> bool:
        """在包含特定文字的段落後添加新段落
        
        Args:
            search_text: 搜尋文字
            new_content: 新內容
            heading_level: 標題層級 (1-9)，None 表示普通段落
            
        Returns:
            bool: 是否找到並添加成功
        """
        if not search_text:
            print(f"{ERROR_SYMBOL} 搜尋文字不能為空")
            return False
            
        found = False
        for i, para in enumerate(self.doc.paragraphs):
            if search_text in para.text:
                # 在找到的段落後插入
                p = para._element
                parent = p.getparent()
                
                # 創建新段落
                if heading_level:
                    new_para = self.doc.add_heading(new_content, level=heading_level)
                else:
                    new_para = self.doc.add_paragraph(new_content)
                
                # 移動到正確位置
                parent.insert(parent.index(p) + 1, new_para._element)
                
                preview = para.text[:50]
                print(f"{SUCCESS_SYMBOL} 已在「{preview}...」後添加內容")
                found = True
                break
        
        if not found:
            print(f"{ERROR_SYMBOL} 找不到包含「{search_text}」的段落")
            
        return found
    
    def replace_text(self, old_text: str, new_text: str, count: int = -1) -> int:
        """替換文字（支援段落和表格）
        
        Args:
            old_text: 要替換的文字
            new_text: 新文字
            count: 替換次數，-1 表示全部替換
            
        Returns:
            int: 實際替換的次數
        """
        if not old_text:
            print(f"{ERROR_SYMBOL} 要替換的文字不能為空")
            return 0
            
        replaced_count = 0
        
        # 替換段落中的文字
        for para in self.doc.paragraphs:
            if old_text in para.text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text, 1 if count > 0 else -1)
                        replaced_count += 1
                        if count > 0 and replaced_count >= count:
                            break
        
        # 替換表格中的文字
        if count < 0 or replaced_count < count:
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old_text in cell.text:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(
                                            old_text, new_text, 1 if count > 0 else -1
                                        )
                                        replaced_count += 1
                                        if count > 0 and replaced_count >= count:
                                            break
        
        if replaced_count > 0:
            print(f"{SUCCESS_SYMBOL} 已替換 {replaced_count} 處「{old_text}」→「{new_text}」")
        else:
            print(f"{ERROR_SYMBOL} 找不到「{old_text}」")
            
        return replaced_count
    
    def delete_paragraph(self, search_text: str) -> bool:
        """刪除包含特定文字的段落
        
        Args:
            search_text: 搜尋文字
            
        Returns:
            bool: 是否找到並刪除成功
        """
        if not search_text:
            print(f"{ERROR_SYMBOL} 搜尋文字不能為空")
            return False
            
        deleted = False
        for para in self.doc.paragraphs:
            if search_text in para.text:
                p = para._element
                p.getparent().remove(p)
                print(f"{SUCCESS_SYMBOL} 已刪除段落: {para.text[:50]}")
                deleted = True
                break
        
        if not deleted:
            print(f"{ERROR_SYMBOL} 找不到包含「{search_text}」的段落")
            
        return deleted
    
    def insert_after_heading(
        self, 
        heading_text: str, 
        content: str, 
        is_heading: bool = False, 
        heading_level: int = DEFAULT_HEADING_LEVEL
    ) -> bool:
        """在特定標題後插入內容
        
        Args:
            heading_text: 標題文字
            content: 要插入的內容
            is_heading: 插入的內容是否為標題
            heading_level: 標題層級 (1-9)
            
        Returns:
            bool: 是否找到並插入成功
        """
        if not heading_text:
            print(f"{ERROR_SYMBOL} 標題文字不能為空")
            return False
            
        found = False
        for i, para in enumerate(self.doc.paragraphs):
            if para.style.name.startswith('Heading') and heading_text in para.text:
                # 找到標題，在它後面插入
                p = para._element
                parent = p.getparent()
                
                if is_heading:
                    new_para = self.doc.add_heading(content, level=heading_level)
                else:
                    new_para = self.doc.add_paragraph(content)
                
                parent.insert(parent.index(p) + 1, new_para._element)
                
                print(f"{SUCCESS_SYMBOL} 已在標題「{para.text}」後插入內容")
                found = True
                break
        
        if not found:
            print(f"{ERROR_SYMBOL} 找不到標題「{heading_text}」")
            
        return found
    
    def add_bullet_points(self, heading_text: str, bullet_points: List[str]) -> bool:
        """在特定標題後添加多個項目符號
        
        Args:
            heading_text: 標題文字
            bullet_points: 項目列表
            
        Returns:
            bool: 是否找到並添加成功
        """
        if not heading_text:
            print(f"{ERROR_SYMBOL} 標題文字不能為空")
            return False
            
        if not bullet_points:
            print(f"{ERROR_SYMBOL} 項目列表不能為空")
            return False
            
        found = False
        for i, para in enumerate(self.doc.paragraphs):
            if para.style.name.startswith('Heading') and heading_text in para.text:
                p = para._element
                parent = p.getparent()
                insert_pos = parent.index(p) + 1
                
                for bullet in bullet_points:
                    new_para = self.doc.add_paragraph(f"• {bullet}")
                    parent.insert(insert_pos, new_para._element)
                    insert_pos += 1
                
                print(f"{SUCCESS_SYMBOL} 已在「{para.text}」後添加 {len(bullet_points)} 個項目")
                found = True
                break
        
        if not found:
            print(f"{ERROR_SYMBOL} 找不到標題「{heading_text}」")
            
        return found

    def add_image(self, image_path: str, width_cm: float = 10.0, position: Optional[str] = None) -> bool:
        """插入圖片
        
        Args:
            image_path: 圖片檔案路徑
            width_cm: 圖片寬度（公分）
            position: 插入位置描述文字，None 表示文檔末尾
            
        Returns:
            bool: 是否成功插入
        """
        if not os.path.exists(image_path):
            print(f"{ERROR_SYMBOL} 圖片檔案不存在: {image_path}")
            return False
        
        try:
            if position:
                # 在特定位置後插入
                for i, para in enumerate(self.doc.paragraphs):
                    if position in para.text:
                        # 在段落後插入新段落並添加圖片
                        p = para._element
                        parent = p.getparent()
                        new_para = self.doc.add_paragraph()
                        parent.insert(parent.index(p) + 1, new_para._element)
                        run = new_para.add_run()
                        run.add_picture(image_path, width=Cm(width_cm))
                        print(f"{SUCCESS_SYMBOL} 已在「{position}」後插入圖片")
                        return True
                print(f"{ERROR_SYMBOL} 找不到位置: {position}")
                return False
            else:
                # 在文檔末尾插入
                para = self.doc.add_paragraph()
                run = para.add_run()
                run.add_picture(image_path, width=Cm(width_cm))
                print(f"{SUCCESS_SYMBOL} 已在文檔末尾插入圖片")
                return True
        except Exception as e:
            print(f"{ERROR_SYMBOL} 插入圖片失敗: {e}")
            return False
    
    def insert_table(self, rows: int, cols: int, data: Optional[List[List[str]]] = None, 
                    position: Optional[str] = None) -> bool:
        """插入表格
        
        Args:
            rows: 行數
            cols: 列數
            data: 表格資料（二維列表）
            position: 插入位置，None 表示文檔末尾
            
        Returns:
            bool: 是否成功插入
        """
        if rows < 1 or cols < 1:
            print(f"{ERROR_SYMBOL} 行列數必須大於 0")
            return False
        
        try:
            if position:
                # 在特定位置後插入
                for para in self.doc.paragraphs:
                    if position in para.text:
                        p = para._element
                        parent = p.getparent()
                        table = self.doc.add_table(rows, cols)
                        parent.insert(parent.index(p) + 1, table._element)
                        
                        # 填充資料
                        if data:
                            for i, row_data in enumerate(data[:rows]):
                                for j, cell_data in enumerate(row_data[:cols]):
                                    table.rows[i].cells[j].text = str(cell_data)
                        
                        print(f"{SUCCESS_SYMBOL} 已插入 {rows}x{cols} 表格")
                        return True
                print(f"{ERROR_SYMBOL} 找不到位置: {position}")
                return False
            else:
                # 在文檔末尾插入
                table = self.doc.add_table(rows, cols)
                
                # 填充資料
                if data:
                    for i, row_data in enumerate(data[:rows]):
                        for j, cell_data in enumerate(row_data[:cols]):
                            table.rows[i].cells[j].text = str(cell_data)
                
                print(f"{SUCCESS_SYMBOL} 已在文檔末尾插入 {rows}x{cols} 表格")
                return True
        except Exception as e:
            print(f"{ERROR_SYMBOL} 插入表格失敗: {e}")
            return False
    
    def update_table_cell(self, table_index: int, row: int, col: int, text: str) -> bool:
        """更新表格儲存格
        
        Args:
            table_index: 表格索引（從 0 開始）
            row: 行索引（從 0 開始）
            col: 列索引（從 0 開始）
            text: 新文字
            
        Returns:
            bool: 是否成功更新
        """
        try:
            if table_index >= len(self.doc.tables):
                print(f"{ERROR_SYMBOL} 表格索引超出範圍（共 {len(self.doc.tables)} 個表格）")
                return False
            
            table = self.doc.tables[table_index]
            
            if row >= len(table.rows):
                print(f"{ERROR_SYMBOL} 行索引超出範圍（共 {len(table.rows)} 行）")
                return False
            
            if col >= len(table.columns):
                print(f"{ERROR_SYMBOL} 列索引超出範圍（共 {len(table.columns)} 列）")
                return False
            
            old_text = table.rows[row].cells[col].text
            table.rows[row].cells[col].text = text
            print(f"{SUCCESS_SYMBOL} 已更新表格[{table_index}][{row},{col}]")
            print(f"  舊值: {old_text}")
            print(f"  新值: {text}")
            return True
        except Exception as e:
            print(f"{ERROR_SYMBOL} 更新表格失敗: {e}")
            return False
    
    def set_paragraph_format(self, search_text: str, font_size: int = 11, 
                            bold: bool = False, italic: bool = False,
                            alignment: Optional[str] = None) -> bool:
        """設定段落格式
        
        Args:
            search_text: 搜尋文字
            font_size: 字體大小
            bold: 是否粗體
            italic: 是否斜體
            alignment: 對齊方式 ('left', 'center', 'right', 'justify')
            
        Returns:
            bool: 是否找到並設定成功
        """
        if not search_text:
            print(f"{ERROR_SYMBOL} 搜尋文字不能為空")
            return False
        
        found = False
        for para in self.doc.paragraphs:
            if search_text in para.text:
                for run in para.runs:
                    run.font.size = Pt(font_size)
                    run.font.bold = bold
                    run.font.italic = italic
                
                # 設定對齊
                if alignment:
                    alignment_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if alignment in alignment_map:
                        para.alignment = alignment_map[alignment]
                
                print(f"{SUCCESS_SYMBOL} 已設定段落格式: {search_text[:50]}")
                found = True
                break
        
        if not found:
            print(f"{ERROR_SYMBOL} 找不到包含「{search_text}」的段落")
        
        return found
    
    def add_page_break(self, after_text: Optional[str] = None) -> bool:
        """插入分頁符號
        
        Args:
            after_text: 在包含此文字的段落後插入，None 表示文檔末尾
            
        Returns:
            bool: 是否成功插入
        """
        try:
            if after_text:
                for para in self.doc.paragraphs:
                    if after_text in para.text:
                        # 在段落後插入分頁
                        p = para._element
                        parent = p.getparent()
                        new_para = self.doc.add_paragraph()
                        parent.insert(parent.index(p) + 1, new_para._element)
                        new_para.add_run().add_break(type=6)  # Page break
                        print(f"{SUCCESS_SYMBOL} 已在「{after_text}」後插入分頁符號")
                        return True
                print(f"{ERROR_SYMBOL} 找不到包含「{after_text}」的段落")
                return False
            else:
                # 在文檔末尾插入
                para = self.doc.add_paragraph()
                para.add_run().add_break(type=6)
                print(f"{SUCCESS_SYMBOL} 已在文檔末尾插入分頁符號")
                return True
        except Exception as e:
            print(f"{ERROR_SYMBOL} 插入分頁符號失敗: {e}")
            return False


def main() -> None:
    """主函數"""
    parser = argparse.ArgumentParser(description='Word 文檔互動式編輯器')
    parser.add_argument('file', help='Word 文檔路徑')
    parser.add_argument('--output', '-o', help='輸出文件路徑（不指定則覆蓋原文件）')
    
    subparsers = parser.add_subparsers(dest='command', help='編輯命令')
    
    # list: 列出文檔結構
    subparsers.add_parser('list', help='列出文檔結構')
    
    # replace: 替換文字
    replace_parser = subparsers.add_parser('replace', help='替換文字')
    replace_parser.add_argument('old', help='要替換的文字')
    replace_parser.add_argument('new', help='新文字')
    replace_parser.add_argument('--count', type=int, default=-1, help='替換次數（-1表示全部）')
    
    # add-after: 在段落後添加內容
    add_parser = subparsers.add_parser('add-after', help='在特定段落後添加內容')
    add_parser.add_argument('search', help='搜尋文字')
    add_parser.add_argument('content', help='要添加的內容')
    add_parser.add_argument('--heading', type=int, help='作為標題（指定層級1-3）')
    
    # insert-after-heading: 在標題後插入
    insert_parser = subparsers.add_parser('insert-after-heading', help='在標題後插入內容')
    insert_parser.add_argument('heading', help='標題文字')
    insert_parser.add_argument('content', help='要插入的內容')
    insert_parser.add_argument('--heading-level', type=int, default=DEFAULT_HEADING_LEVEL, 
                              help='作為標題層級')
    insert_parser.add_argument('--is-heading', action='store_true', help='插入的內容是標題')
    
    # delete: 刪除段落
    delete_parser = subparsers.add_parser('delete', help='刪除段落')
    delete_parser.add_argument('search', help='要刪除的段落（搜尋文字）')
    
    # add-bullets: 添加項目符號
    bullets_parser = subparsers.add_parser('add-bullets', help='在標題後添加項目符號')
    bullets_parser.add_argument('heading', help='標題文字')
    bullets_parser.add_argument('bullets', nargs='+', help='項目內容（可多個）')
    
    args = parser.parse_args()
    
    if not args.command:
        parser.print_help()
        return
    
    # 載入文檔
    try:
        editor = WordEditor(args.file)
    except (FileNotFoundError, ValueError, RuntimeError) as e:
        print(f"{ERROR_SYMBOL} {e}")
        sys.exit(1)
    
    # 執行命令
    try:
        if args.command == 'list':
            editor.list_structure()
            return
        
        elif args.command == 'replace':
            editor.replace_text(args.old, args.new, args.count)
        
        elif args.command == 'add-after':
            editor.add_paragraph_after(args.search, args.content, args.heading)
        
        elif args.command == 'insert-after-heading':
            editor.insert_after_heading(args.heading, args.content, 
                                       args.is_heading, args.heading_level)
        
        elif args.command == 'delete':
            editor.delete_paragraph(args.search)
        
        elif args.command == 'add-bullets':
            editor.add_bullet_points(args.heading, args.bullets)
        
        # 儲存
        editor.save(args.output)
        
    except Exception as e:
        print(f"{ERROR_SYMBOL} 操作失敗: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
