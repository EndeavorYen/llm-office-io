from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import sys

def set_font(run, chinese_font="微軟正黑體", english_font="Aptos"):
    """設定中英文字體"""
    run.font.name = english_font
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)

def set_paragraph_font(paragraph, chinese_font="微軟正黑體", english_font="Aptos"):
    """設定段落中所有文字的字體"""
    for run in paragraph.runs:
        set_font(run, chinese_font, english_font)

def create_restructured_document(input_file, output_file):
    """創建重新結構化的專業文檔"""
    
    # 讀取原始文檔以獲取表格數據
    original_doc = Document(input_file)
    
    # 創建新文檔
    doc = Document()
    
    # 設定預設字體
    style = doc.styles['Normal']
    style.font.name = 'Aptos'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    
    # ============ 封面頁 ============
    # 添加標題
    title = doc.add_heading('Smart Cell 教育訓練課程', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # 深藍色
    
    doc.add_paragraph()  # 空行
    
    # 副標題
    subtitle = doc.add_paragraph('完整培訓指南')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 課程信息框
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        '📅 課程日期：2025/12/03 - 2025/12/05\n'
        '👨‍🏫 講師：Tim (資深工程師)\n'
        '📧 Email: tim@example.com'
    )
    info_run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 版本信息
    version = doc.add_paragraph('Version 1.0 - 2025年12月')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.runs[0].font.size = Pt(10)
    version.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # 分頁
    doc.add_page_break()
    
    # ============ 目錄 ============
    toc_heading = doc.add_heading('📑 目錄', level=1)
    toc_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    toc_items = [
        ('一、課程概覽', '3'),
        ('    1.1 課程時程表', '3'),
        ('    1.2 課程目標', '4'),
        ('    1.3 學習重點', '4'),
        ('二、Smart Cell 系統介紹', '5'),
        ('    2.1 系統概述', '5'),
        ('    2.2 系統特色', '5'),
        ('三、課前準備', '6'),
        ('    3.1 預備知識', '6'),
        ('    3.2 環境設定', '6'),
        ('    3.3 準備事項', '7'),
        ('四、課程詳細資訊', '7'),
        ('    4.1 課程形式', '7'),
        ('    4.2 教材提供', '7'),
        ('    4.3 注意事項', '8'),
        ('五、預期成果與聯絡方式', '8'),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item + ' ')
        
        # 添加點線
        dots = '.' * (60 - len(item) - len(page))
        p.add_run(dots + ' ')
        p.add_run(page)
        p.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    # 分頁
    doc.add_page_break()
    
    # ============ 第一部分：課程概覽 ============
    section1 = doc.add_heading('一、課程概覽', level=1)
    section1.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_heading('1.1 課程時程表', level=2)
    
    # Day 1 表格
    doc.add_heading('📌 Day 1: Smart Cell 系統架構介紹', level=3)
    p = doc.add_paragraph()
    p.add_run('📅 日期：').bold = True
    p.add_run('2025/12/03\n')
    p.add_run('⏰ 時間：').bold = True
    p.add_run('13:30-15:30\n')
    p.add_run('📍 地點：').bold = True
    p.add_run('Teams 線上會議')
    
    # 複製原始表格（如果存在）
    if len(original_doc.tables) > 0:
        table1 = original_doc.tables[0]
        new_table1 = doc.add_table(rows=len(table1.rows), cols=len(table1.columns))
        new_table1.style = 'Light Grid Accent 1'
        
        for i, row in enumerate(table1.rows):
            for j, cell in enumerate(row.cells):
                new_table1.rows[i].cells[j].text = cell.text
    
    doc.add_paragraph()
    
    # Day 2 表格
    doc.add_heading('📌 Day 2: Smart Cell 使用方法', level=3)
    p = doc.add_paragraph()
    p.add_run('📅 日期：').bold = True
    p.add_run('2025/12/04\n')
    p.add_run('⏰ 時間：').bold = True
    p.add_run('13:30-15:30\n')
    p.add_run('📍 地點：').bold = True
    p.add_run('Teams 線上會議')
    
    if len(original_doc.tables) > 1:
        table2 = original_doc.tables[1]
        new_table2 = doc.add_table(rows=len(table2.rows), cols=len(table2.columns))
        new_table2.style = 'Light Grid Accent 1'
        
        for i, row in enumerate(table2.rows):
            for j, cell in enumerate(row.cells):
                new_table2.rows[i].cells[j].text = cell.text
    
    doc.add_paragraph()
    
    # Day 3 表格
    doc.add_heading('📌 Day 3: Smart Cell 實際操作', level=3)
    p = doc.add_paragraph()
    p.add_run('📅 日期：').bold = True
    p.add_run('2025/12/05\n')
    p.add_run('⏰ 時間：').bold = True
    p.add_run('13:00-17:00\n')
    p.add_run('📍 地點：').bold = True
    p.add_run('新竹研發中心 17F 1707實驗室')
    
    if len(original_doc.tables) > 2:
        table3 = original_doc.tables[2]
        new_table3 = doc.add_table(rows=len(table3.rows), cols=len(table3.columns))
        new_table3.style = 'Light Grid Accent 1'
        
        for i, row in enumerate(table3.rows):
            for j, cell in enumerate(row.cells):
                new_table3.rows[i].cells[j].text = cell.text
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # 1.2 課程目標
    doc.add_heading('1.2 課程目標', level=2)
    
    goal_para = doc.add_paragraph(
        '本次教育訓練旨在讓學員全面了解 Smart Cell 系統的架構、使用方法和開發流程。'
        '透過三天循序漸進的課程安排，從理論到實作，幫助學員快速上手並具備獨立開發與測試的能力。'
    )
    goal_para.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # 1.3 學習重點
    doc.add_heading('1.3 學習重點', level=2)
    
    doc.add_heading('Day 1 重點：系統架構與設計', level=3)
    doc.add_paragraph('✓ 理解 Smart Cell 的整體程式架構與模組關係')
    doc.add_paragraph('✓ 掌握系統需求與設計規格的核心概念')
    doc.add_paragraph('✓ 了解資料儲存格式與數據流向')
    
    doc.add_heading('Day 2 重點：工具使用與流程', level=3)
    doc.add_paragraph('✓ 學會使用 Smart Cell CLI 進行系統操作')
    doc.add_paragraph('✓ 熟悉自動化測試的配置與執行方法')
    doc.add_paragraph('✓ 了解開發環境建置、編譯流程與 OTA 更新機制')
    doc.add_paragraph('✓ 掌握 Git Repository 的使用與協作流程')
    
    doc.add_heading('Day 3 重點：實機操作與整合', level=3)
    doc.add_paragraph('✓ 實際操作 BMS CLI 的各項功能')
    doc.add_paragraph('✓ 實際操作 Module Controller CLI')
    doc.add_paragraph('✓ 實際操作 Smart Cell CLI 進行系統調試')
    doc.add_paragraph('✓ 使用自動化測試平台 (ATE) 執行完整測試流程')
    
    doc.add_page_break()
    
    # ============ 第二部分：Smart Cell 系統介紹 ============
    section2 = doc.add_heading('二、Smart Cell 系統介紹', level=1)
    section2.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_heading('2.1 系統概述', level=2)
    
    overview_para = doc.add_paragraph(
        'Smart Cell 是一個先進的電池管理系統 (Battery Management System, BMS)，'
        '整合了智能控制、數據採集和自動化測試功能。系統架構包含三個主要組件：'
    )
    overview_para.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # 系統組件
    p1 = doc.add_paragraph()
    p1.add_run('🔹 BMS (Battery Management System)\n').bold = True
    p1.add_run('   負責電池組的監控、保護和平衡管理，確保電池系統的安全運行。')
    
    p2 = doc.add_paragraph()
    p2.add_run('🔹 Module Controller\n').bold = True
    p2.add_run('   控制各個模組的運作，處理模組間的通訊與協調。')
    
    p3 = doc.add_paragraph()
    p3.add_run('🔹 Smart Cell CLI\n').bold = True
    p3.add_run('   提供命令列介面，方便開發人員進行系統配置、監控和調試。')
    
    doc.add_paragraph()
    
    doc.add_heading('2.2 系統主要特色', level=2)
    
    doc.add_paragraph('⚡ 即時監控：提供電池狀態的即時數據採集與分析')
    doc.add_paragraph('🛡️ 智能保護：多層次安全保護機制，防止過充、過放、過溫等異常狀況')
    doc.add_paragraph('🤖 自動化測試：整合 ATE (Automated Test Equipment) 平台，提升測試效率')
    doc.add_paragraph('📦 彈性擴展：模組化設計，支援不同規模的電池系統配置')
    doc.add_paragraph('🔄 OTA 更新：支援遠端韌體更新，降低維護成本')
    
    doc.add_page_break()
    
    # ============ 第三部分：課前準備 ============
    section3 = doc.add_heading('三、課前準備', level=1)
    section3.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_heading('3.1 預備知識', level=2)
    
    doc.add_paragraph('為了確保學習效果，建議學員具備以下基礎知識：')
    
    doc.add_heading('必備技能 ⭐', level=3)
    doc.add_paragraph('• 基礎程式設計能力（C/C++ 或 Python）')
    doc.add_paragraph('• 命令列介面 (CLI) 基本操作經驗')
    doc.add_paragraph('• Git 版本控制系統基本概念')
    doc.add_paragraph('• Linux/Unix 環境基本操作')
    
    doc.add_heading('加分技能 ✨', level=3)
    doc.add_paragraph('• 嵌入式系統開發經驗')
    doc.add_paragraph('• 電池管理系統相關知識')
    doc.add_paragraph('• 自動化測試經驗')
    doc.add_paragraph('• 串列通訊協定 (UART, SPI, I2C) 了解')
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 環境設定', level=2)
    
    doc.add_heading('軟體需求 💻', level=3)
    doc.add_paragraph('• 作業系統：Windows 10/11 或 Linux (Ubuntu 20.04+)')
    doc.add_paragraph('• Python 3.8 或更新版本')
    doc.add_paragraph('• Git 版本控制工具')
    doc.add_paragraph('• IDE 或文字編輯器 (VS Code 推薦)')
    doc.add_paragraph('• Microsoft Teams (線上課程使用)')
    
    doc.add_heading('硬體需求（Day 3 實機操作）🔧', level=3)
    doc.add_paragraph('• 筆記型電腦')
    doc.add_paragraph('• USB 連接線')
    doc.add_paragraph('• Smart Cell 開發板（現場提供）')
    
    doc.add_paragraph()
    
    doc.add_heading('3.3 課前準備事項 ✅', level=2)
    
    checklist = [
        '確認 Teams 帳號可正常登入',
        '安裝 Git 並設定基本配置 (user.name, user.email)',
        '安裝 Python 3.8+ 並確認可在命令列執行',
        '準備筆記本或數位筆記工具記錄重點',
        '（Day 3 參加者）確認可到達新竹研發中心 17F'
    ]
    
    for i, item in enumerate(checklist, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    doc.add_page_break()
    
    # ============ 第四部分：課程詳細資訊 ============
    section4 = doc.add_heading('四、課程詳細資訊', level=1)
    section4.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_heading('4.1 課程形式', level=2)
    
    p = doc.add_paragraph()
    p.add_run('🌐 Day 1-2：').bold = True
    p.add_run('線上授課，透過 Teams 進行，包含簡報說明與即時 Q&A\n')
    p.add_run('🏢 Day 3：').bold = True
    p.add_run('實體課程，在新竹研發中心進行實機操作與互動教學')
    
    doc.add_paragraph()
    
    doc.add_heading('4.2 教材提供', level=2)
    
    doc.add_paragraph('📄 課程簡報檔（PDF 格式）')
    doc.add_paragraph('📖 系統操作手冊')
    doc.add_paragraph('💾 範例程式碼與測試腳本')
    doc.add_paragraph('🔑 Git Repository 存取權限')
    
    doc.add_paragraph()
    
    doc.add_heading('4.3 注意事項 ⚠️', level=2)
    
    doc.add_paragraph('⏰ 請準時參加線上會議，遲到可能錯過重要內容')
    doc.add_paragraph('💼 Day 3 實體課程請攜帶筆記型電腦')
    doc.add_paragraph('📚 建議課前預習相關文件，提升學習效率')
    doc.add_paragraph('💬 課程中歡迎提問與討論')
    doc.add_paragraph('⚠️ 實機操作時請小心操作設備，避免損壞')
    
    doc.add_page_break()
    
    # ============ 第五部分：預期成果與聯絡方式 ============
    section5 = doc.add_heading('五、預期成果與聯絡方式', level=1)
    section5.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_heading('預期成果 🎯', level=2)
    
    doc.add_paragraph('完成三天課程後，學員將能夠：')
    doc.add_paragraph()
    
    outcomes = [
        '獨立使用 Smart Cell 系統進行開發與測試',
        '理解系統架構並能進行基本的除錯與問題排查',
        '使用自動化測試平台提升開發效率',
        '掌握完整的開發流程從編譯到 OTA 更新',
        '具備團隊協作開發的基礎能力'
    ]
    
    for i, outcome in enumerate(outcomes, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. ').bold = True
        p.add_run(outcome)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc.add_heading('聯絡資訊 📞', level=2)
    
    contact_box = doc.add_paragraph()
    contact_box.paragraph_format.left_indent = Inches(0.5)
    contact_box.add_run('👨‍🏫 講師：').bold = True
    contact_box.add_run('Tim (資深工程師)\n')
    contact_box.add_run('📧 Email：').bold = True
    contact_box.add_run('tim@example.com\n\n')
    contact_box.add_run('如有任何問題，請於課程前或課程中隨時提出\n')
    contact_box.add_run('Q&A 時段將預留充足時間解答疑問')
    
    # 底部分隔線
    doc.add_paragraph()
    doc.add_paragraph('─' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    thanks = doc.add_paragraph('期待與您在課堂上見面！')
    thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    thanks.runs[0].font.size = Pt(12)
    thanks.runs[0].font.bold = True
    thanks.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    # 應用字體到整個文檔
    apply_fonts_to_document(doc)
    
    # 儲存文檔
    doc.save(output_file)
    print(f"\n✓ 重新結構化的文檔已創建: {output_file}")
    print("✓ 包含封面頁、目錄和清晰的分節結構")
    print("✓ 字體設定：中文 - 微軟正黑體，英文 - Aptos")

def apply_fonts_to_document(doc, chinese_font="微軟正黑體", english_font="Aptos"):
    """應用字體到整個文檔的所有段落和表格"""
    # 設定所有段落
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_font(run, chinese_font, english_font)
    
    # 設定所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_font(run, chinese_font, english_font)

if __name__ == "__main__":
    input_file = r"SmartCell教育訓練時程_完整版.docx"
    output_file = r"SmartCell教育訓練課程_專業版.docx"
    
    try:
        create_restructured_document(input_file, output_file)
    except Exception as e:
        print(f"錯誤: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
