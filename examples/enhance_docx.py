from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

def enhance_document(input_file, output_file):
    """Add detailed sections to the training document"""
    
    # Load existing document
    doc = Document(input_file)
    
    # Add page break before new content
    doc.add_page_break()
    
    # Section 1: Smart Cell 系統概述
    heading1 = doc.add_heading('Smart Cell 系統概述', level=1)
    heading1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph(
        'Smart Cell 是一個先進的電池管理系統(Battery Management System, BMS)，'
        '整合了智能控制、數據採集和自動化測試功能。系統架構包含三個主要組件：'
    )
    
    # Add system components
    p1 = doc.add_paragraph()
    p1.add_run('• BMS (Battery Management System)\n').bold = True
    p1.add_run('  負責電池組的監控、保護和平衡管理，確保電池系統的安全運行。')
    
    p2 = doc.add_paragraph()
    p2.add_run('• Module Controller\n').bold = True
    p2.add_run('  控制各個模組的運作，處理模組間的通訊與協調。')
    
    p3 = doc.add_paragraph()
    p3.add_run('• Smart Cell CLI\n').bold = True
    p3.add_run('  提供命令列介面，方便開發人員進行系統配置、監控和調試。')
    
    doc.add_paragraph()  # Empty line
    
    # System features
    doc.add_heading('系統主要特色', level=2)
    doc.add_paragraph('• 即時監控：提供電池狀態的即時數據採集與分析')
    doc.add_paragraph('• 智能保護：多層次安全保護機制，防止過充、過放、過溫等異常狀況')
    doc.add_paragraph('• 自動化測試：整合 ATE (Automated Test Equipment) 平台，提升測試效率')
    doc.add_paragraph('• 彈性擴展：模組化設計，支援不同規模的電池系統配置')
    doc.add_paragraph('• OTA 更新：支援遠端韌體更新，降低維護成本')
    
    doc.add_paragraph()
    
    # Section 2: 課程目標和學習重點
    doc.add_heading('課程目標和學習重點', level=1)
    
    doc.add_heading('課程目標', level=2)
    doc.add_paragraph(
        '本次教育訓練旨在讓學員全面了解 Smart Cell 系統的架構、使用方法和開發流程。'
        '透過三天循序漸進的課程安排，從理論到實作，幫助學員快速上手並具備獨立開發與測試的能力。'
    )
    
    doc.add_paragraph()
    
    doc.add_heading('學習重點', level=2)
    
    doc.add_heading('Day 1 重點：系統架構與設計', level=3)
    doc.add_paragraph('• 理解 Smart Cell 的整體程式架構與模組關係')
    doc.add_paragraph('• 掌握系統需求與設計規格的核心概念')
    doc.add_paragraph('• 了解資料儲存格式與數據流向')
    
    doc.add_heading('Day 2 重點：工具使用與流程', level=3)
    doc.add_paragraph('• 學會使用 Smart Cell CLI 進行系統操作')
    doc.add_paragraph('• 熟悉自動化測試的配置與執行方法')
    doc.add_paragraph('• 了解開發環境建置、編譯流程與 OTA 更新機制')
    doc.add_paragraph('• 掌握 Git Repository 的使用與協作流程')
    
    doc.add_heading('Day 3 重點：實機操作與整合', level=3)
    doc.add_paragraph('• 實際操作 BMS CLI 的各項功能')
    doc.add_paragraph('• 實際操作 Module Controller CLI')
    doc.add_paragraph('• 實際操作 Smart Cell CLI 進行系統調試')
    doc.add_paragraph('• 使用自動化測試平台 (ATE) 執行完整測試流程')
    
    doc.add_paragraph()
    
    # Add page break
    doc.add_page_break()
    
    # Section 3: 預備知識和環境設定
    doc.add_heading('預備知識和環境設定', level=1)
    
    doc.add_heading('預備知識', level=2)
    doc.add_paragraph(
        '為了確保學習效果，建議學員具備以下基礎知識：'
    )
    
    doc.add_heading('必備技能', level=3)
    doc.add_paragraph('• 基礎程式設計能力（C/C++ 或 Python）')
    doc.add_paragraph('• 命令列介面 (CLI) 基本操作經驗')
    doc.add_paragraph('• Git 版本控制系統基本概念')
    doc.add_paragraph('• Linux/Unix 環境基本操作')
    
    doc.add_heading('加分技能', level=3)
    doc.add_paragraph('• 嵌入式系統開發經驗')
    doc.add_paragraph('• 電池管理系統相關知識')
    doc.add_paragraph('• 自動化測試經驗')
    doc.add_paragraph('• 串列通訊協定 (UART, SPI, I2C) 了解')
    
    doc.add_paragraph()
    
    doc.add_heading('環境設定', level=2)
    
    doc.add_heading('軟體需求', level=3)
    doc.add_paragraph('• 作業系統：Windows 10/11 或 Linux (Ubuntu 20.04+)')
    doc.add_paragraph('• Python 3.8 或更新版本')
    doc.add_paragraph('• Git 版本控制工具')
    doc.add_paragraph('• IDE 或文字編輯器 (VS Code 推薦)')
    doc.add_paragraph('• Microsoft Teams (線上課程使用)')
    
    doc.add_heading('硬體需求（Day 3 實機操作）', level=3)
    doc.add_paragraph('• 筆記型電腦')
    doc.add_paragraph('• USB 連接線')
    doc.add_paragraph('• Smart Cell 開發板（現場提供）')
    
    doc.add_heading('課前準備事項', level=3)
    doc.add_paragraph('1. 確認 Teams 帳號可正常登入')
    doc.add_paragraph('2. 安裝 Git 並設定基本配置 (user.name, user.email)')
    doc.add_paragraph('3. 安裝 Python 3.8+ 並確認可在命令列執行')
    doc.add_paragraph('4. 準備筆記本或數位筆記工具記錄重點')
    doc.add_paragraph('5. （Day 3 參加者）確認可到達新竹研發中心 17F')
    
    doc.add_paragraph()
    
    # Section 4: 其他重要資訊
    doc.add_heading('其他重要資訊', level=1)
    
    doc.add_heading('課程形式', level=2)
    doc.add_paragraph('• Day 1-2：線上授課，透過 Teams 進行，包含簡報說明與即時 Q&A')
    doc.add_paragraph('• Day 3：實體課程，在新竹研發中心進行實機操作與互動教學')
    
    doc.add_paragraph()
    
    doc.add_heading('教材提供', level=2)
    doc.add_paragraph('• 課程簡報檔（PDF 格式）')
    doc.add_paragraph('• 系統操作手冊')
    doc.add_paragraph('• 範例程式碼與測試腳本')
    doc.add_paragraph('• Git Repository 存取權限')
    
    doc.add_paragraph()
    
    doc.add_heading('聯絡資訊', level=2)
    doc.add_paragraph('講師：Tim')
    doc.add_paragraph('如有任何問題，請於課程前或課程中隨時提出')
    doc.add_paragraph('Q&A 時段將預留充足時間解答疑問')
    
    doc.add_paragraph()
    
    doc.add_heading('注意事項', level=2)
    doc.add_paragraph('• 請準時參加線上會議，遲到可能錯過重要內容')
    doc.add_paragraph('• Day 3 實體課程請攜帶筆記型電腦')
    doc.add_paragraph('• 建議課前預習相關文件，提升學習效率')
    doc.add_paragraph('• 課程中歡迎提問與討論')
    doc.add_paragraph('• 實機操作時請小心操作設備，避免損壞')
    
    doc.add_paragraph()
    
    doc.add_heading('預期成果', level=2)
    doc.add_paragraph(
        '完成三天課程後，學員將能夠：'
    )
    doc.add_paragraph('1. 獨立使用 Smart Cell 系統進行開發與測試')
    doc.add_paragraph('2. 理解系統架構並能進行基本的除錯與問題排查')
    doc.add_paragraph('3. 使用自動化測試平台提升開發效率')
    doc.add_paragraph('4. 掌握完整的開發流程從編譯到 OTA 更新')
    doc.add_paragraph('5. 具備團隊協作開發的基礎能力')
    
    # Save the enhanced document
    doc.save(output_file)
    print(f"Enhanced document saved to: {output_file}")

if __name__ == "__main__":
    input_file = r"SmartCell教育訓練時程.docx"
    output_file = r"SmartCell教育訓練時程_完整版.docx"
    
    try:
        enhance_document(input_file, output_file)
        print("Document enhancement completed successfully!")
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)
